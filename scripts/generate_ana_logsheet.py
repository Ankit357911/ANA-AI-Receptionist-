from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


TEAM_MEMBERS = "Ankit Chandra Karn, Srijan Basnet, Suraj Panthi"
OUTPUT_PATH = "ANA_Internship_Logsheet.docx"


ROWS = [
    ("08-May-2026", "Discussed the AI receptionist project idea and identified the main requirements for a college enquiry system.", "Collect college information and prepare initial project structure."),
    ("12-May-2026", "Collected basic college data and created folders for frontend, backend, data, and AI service modules.", "Set up Python environment and install required libraries."),
    ("13-May-2026", "Configured Python virtual environment and installed FastAPI, FAISS, sentence-transformers, OpenAI client, and supporting packages.", "Study RAG workflow and plan document indexing."),
    ("14-May-2026", "Prepared text documents for college overview, admission, programs, facilities, and FAQs.", "Create chunking and embedding pipeline."),
    ("15-May-2026", "Developed document chunking and embedding logic for college information files.", "Build FAISS vector index for retrieval."),
    ("19-May-2026", "Created FAISS indexes for college documents and tested basic similarity search.", "Connect retrieval results with answer generation."),
    ("20-May-2026", "Implemented retriever and router logic to search across multiple document indexes.", "Improve relevance threshold and source selection."),
    ("21-May-2026", "Connected retrieved context with local LLM response generation for college questions.", "Create command-line chatbot for testing."),
    ("22-May-2026", "Built a basic CLI receptionist chatbot and tested common questions about courses and admission.", "Develop FastAPI backend for web usage."),
    ("26-May-2026", "Created FastAPI chat endpoint and session-based conversation handling.", "Design simple browser interface for the receptionist."),
    ("27-May-2026", "Developed frontend chat interface with message display, input box, and backend API connection.", "Add voice input support."),
    ("28-May-2026", "Implemented speech-to-text service using Faster Whisper and tested microphone transcription.", "Improve transcription accuracy for college terms."),
    ("29-May-2026", "Added post-processing rules for terms like KCC, BCA-IT, BCAIT, and Ravi sir.", "Integrate text-to-speech output."),
    ("02-Jun-2026", "Integrated Kokoro text-to-speech service for spoken receptionist replies.", "Create voice mode workflow in the frontend."),
    ("03-Jun-2026", "Added voice recording, silence detection, transcription request, and spoken response playback in the UI.", "Test complete voice conversation flow."),
    ("04-Jun-2026", "Tested voice mode with college questions and fixed issues in microphone recording and response playback.", "Improve answer formatting and remove markdown output."),
    ("05-Jun-2026", "Updated system prompt and response cleaning to keep answers plain, polite, and receptionist-like.", "Add fallback handling for missing document answers."),
    ("09-Jun-2026", "Added fallback answer for BCA-IT queries and improved query normalization.", "Work on AI receptionist visual interface."),
    ("10-Jun-2026", "Designed a live receptionist UI with avatar area, status display, microphone, speaker, and prompt controls.", "Start Wav2Lip integration for animated video replies."),
    ("11-Jun-2026", "Added Wav2Lip service wrapper and configured avatar image, checkpoint path, runtime media, and ffmpeg support.", "Connect TTS audio with Wav2Lip video generation."),
    ("12-Jun-2026", "Connected generated speech audio with Wav2Lip inference and returned video responses through FastAPI.", "Add video playback to receptionist frontend."),
    ("23-Jun-2026", "Integrated synced avatar video playback in the live receptionist interface.", "Improve caching for repeated video responses."),
    ("24-Jun-2026", "Added Wav2Lip video cache based on response text to avoid regenerating repeated replies.", "Test cached and uncached receptionist responses."),
    ("25-Jun-2026", "Tested receptionist responses, video playback, audio output, and cache behavior with different questions.", "Improve persistent cache storage."),
    ("26-Jun-2026", "Moved Wav2Lip cache to a permanent cache folder and saved generated video, audio, and text metadata.", "Add schedule awareness for teacher availability questions."),
    ("30-Jun-2026", "Implemented time-aware schedule handling for Ravi sir using the daily schedule document and Nepal time.", "Generalize schedule handling for more teachers."),
    ("01-Jul-2026", "Improved schedule service to support teacher schedule queries and current class/activity lookup.", "Fix greeting and conversation context issues."),
    ("02-Jul-2026", "Added deterministic greeting handling so simple inputs like hey and hello do not continue old topics.", "Optimize response and video cache usage."),
    ("03-Jul-2026", "Added response cache for stable questions such as BCA-IT course enquiries to reuse previous answer videos.", "Perform final testing of chat, voice, schedule, and Wav2Lip modules."),
    ("07-Jul-2026", "Completed final testing, fixed major issues, prepared project documentation, and reviewed the complete AI receptionist workflow.", "Prepare final submission and presentation."),
]


def set_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)


def main():
    document = Document()
    title = document.add_paragraph("INTERNSHIP-LOGSHEET")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    members = document.add_paragraph()
    members.add_run("Group Members : ").bold = True
    members.add_run(TEAM_MEMBERS)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Date", "Completed Task", "To Do List", "Supervisor Signature"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for date, completed, todo in ROWS:
        cells = table.add_row().cells
        set_cell_text(cells[0], date)
        set_cell_text(cells[1], completed)
        set_cell_text(cells[2], todo)
        set_cell_text(cells[3], "")

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
