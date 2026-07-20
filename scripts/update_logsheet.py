import docx
from docx.shared import Pt
import os
import sys

# Add project root to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from scripts.generate_ana_logsheet import ROWS

def update_docx_file(filepath):
    if not os.path.exists(filepath):
        print(f"Error: {filepath} does not exist.")
        return False
        
    print(f"Updating {filepath}...")
    doc = docx.Document(filepath)
    
    # 1. Update group members paragraph
    members_p = None
    for p in doc.paragraphs:
        if "Group Members" in p.text:
            members_p = p
            break
            
    if members_p:
        # Clear existing text/runs
        members_p.text = ""
        r1 = members_p.add_run("Group Members : ")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        r1.bold = True
        
        r2 = members_p.add_run("Ankit Chandra Karn, Srijan Basnet, Suraj Panthi")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
    else:
        print("Warning: Group Members paragraph not found!")
        
    # 2. Update table cells
    if doc.tables:
        table = doc.tables[0]
        
        # Ensure table rows count is sufficient
        for i, (date, completed, todo) in enumerate(ROWS):
            if i + 1 < len(table.rows):
                row = table.rows[i + 1]
                
                # Check date
                doc_date = row.cells[0].text.strip()
                if doc_date != date:
                    print(f"Date mismatch at row {i+1}: doc='{doc_date}', py='{date}'")
                    row.cells[0].text = ""
                    r_date = row.cells[0].paragraphs[0].add_run(date)
                    r_date.font.name = "Times New Roman"
                    r_date.font.size = Pt(11)
                    
                # Update Completed Task cell
                row.cells[1].text = ""
                r_comp = row.cells[1].paragraphs[0].add_run(completed)
                r_comp.font.name = "Times New Roman"
                r_comp.font.size = Pt(11)
                
                # Update To Do List cell
                row.cells[2].text = ""
                r_todo = row.cells[2].paragraphs[0].add_run(todo)
                r_todo.font.name = "Times New Roman"
                r_todo.font.size = Pt(11)
                
                # Clear signature cell
                row.cells[3].text = ""
            else:
                print(f"Warning: Table does not have row {i+1}")
    else:
        print("Warning: No tables found in document!")
        
    doc.save(filepath)
    print(f"Successfully saved {filepath}\n")
    return True

if __name__ == "__main__":
    paths = [
        "C:/Users/Acer/Desktop/ml_chatbot/Date.docx",
        "C:/Users/Acer/Downloads/Date.docx"
    ]
    for p in paths:
        update_docx_file(p)
